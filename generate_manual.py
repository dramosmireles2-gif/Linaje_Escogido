#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generador del Manual de Código — Linaje Escogido
Produce manual-codigo.docx con documentación completa del proyecto.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ──────────────────────────────────────────────
# HELPERS
# ──────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Aplica color de fondo a una celda de tabla."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_para_shading(para, hex_color: str):
    """Aplica sombreado de párrafo (fondo) para bloques de código."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def add_page_number(doc):
    """Agrega número de página centrado en el pie de cada página."""
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def add_heading(doc, text: str, level: int = 1):
    """Agrega encabezado con estilo personalizado."""
    para = doc.add_paragraph(style=f'Heading {level}')
    run = para.add_run(text)
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2A, 0x2A, 0x6A)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return para


def add_body(doc, text: str, bold_parts: list = None):
    """Agrega párrafo de texto normal."""
    para = doc.add_paragraph(style='Normal')
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.space_before = Pt(2)
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    return para


def add_note(doc, text: str, tipo: str = 'nota'):
    """Agrega un bloque de nota/aviso destacado."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.right_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    icons = {'nota': '📝 Nota:', 'tip': '💡 Consejo:', 'aviso': '⚠️  Aviso:'}
    label = icons.get(tipo, 'ℹ️  Info:')
    run_label = para.add_run(label + ' ')
    run_label.font.bold = True
    run_label.font.name = 'Calibri'
    run_label.font.size = Pt(10)
    run_text = para.add_run(text)
    run_text.font.name = 'Calibri'
    run_text.font.size = Pt(10)
    colors = {'nota': 'EAF4FF', 'tip': 'EAFFF0', 'aviso': 'FFF8EA'}
    set_para_shading(para, colors.get(tipo, 'F0F0F0'))
    return para


def add_code(doc, code_text: str, lang: str = ''):
    """Agrega bloque de código con fondo gris, fuente Courier New."""
    lines = code_text.strip('\n').split('\n')
    for i, line in enumerate(lines):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0) if i > 0 else Pt(6)
        para.paragraph_format.space_after = Pt(0) if i < len(lines) - 1 else Pt(6)
        para.paragraph_format.left_indent = Cm(0.5)
        para.paragraph_format.right_indent = Cm(0.5)
        run = para.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5E)
        set_para_shading(para, 'F2F2F2')
    return para


def add_bullet(doc, text: str, level: int = 0):
    """Agrega elemento de lista con viñeta."""
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return para


def add_numbered(doc, text: str):
    """Agrega elemento de lista numerada."""
    para = doc.add_paragraph(style='List Number')
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return para


def make_table(doc, headers: list, rows: list, col_widths: list = None):
    """Crea tabla con encabezado oscuro y filas alternadas."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Cabecera
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_bg(cell, '1A1A1A')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xF5, 0xEF, 0xE6)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Filas
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = 'FFFFFF' if r_idx % 2 == 0 else 'F8F5F1'
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(cell_text)
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Calibri'
            run.font.size = Pt(10)

    # Ancho de columnas
    if col_widths:
        for i, col in enumerate(table.columns):
            col.width = Cm(col_widths[i])

    doc.add_paragraph()
    return table


# ──────────────────────────────────────────────
# DOCUMENTO PRINCIPAL
# ──────────────────────────────────────────────

doc = Document()

# Márgenes de página
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.2)

# Estilos base
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

add_page_number(doc)

# ═══════════════════════════════════════════════════════
# 1. PORTADA
# ═══════════════════════════════════════════════════════

# Espacio superior
for _ in range(6):
    doc.add_paragraph()

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run('Manual de Código')
title_run.font.name = 'Calibri'
title_run.font.size = Pt(36)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle_para.add_run('Linaje Escogido')
sub_run.font.name = 'Calibri'
sub_run.font.size = Pt(28)
sub_run.font.bold = False
sub_run.font.color.rgb = RGBColor(0x2A, 0x2A, 0x8A)

doc.add_paragraph()

desc_para = doc.add_paragraph()
desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
desc_run = desc_para.add_run('Guía para entender y modificar el sitio web')
desc_run.font.name = 'Calibri'
desc_run.font.size = Pt(16)
desc_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()

iglesia_para = doc.add_paragraph()
iglesia_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
iglesia_run = iglesia_para.add_run('Iglesia Linaje Escogido · Reynosa, Tamaulipas')
iglesia_run.font.name = 'Calibri'
iglesia_run.font.size = Pt(13)
iglesia_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_para.add_run('Versión 1.0 — 2026')
date_run.font.name = 'Calibri'
date_run.font.size = Pt(12)
date_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

for _ in range(4):
    doc.add_paragraph()

footer_cover = doc.add_paragraph()
footer_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
fc_run = footer_cover.add_run('Una Iglesia Funcional · Apasionada · Sensible al Espíritu Santo')
fc_run.font.name = 'Calibri'
fc_run.font.size = Pt(10)
fc_run.font.italic = True
fc_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 2. RESUMEN GENERAL
# ═══════════════════════════════════════════════════════

add_heading(doc, '1. Resumen General del Proyecto', 1)
add_body(doc,
    'El sitio web de Linaje Escogido es una página pública de iglesia diseñada para presentar '
    'la comunidad, sus servicios, anuncios, pastores y formularios de contacto. Está construida '
    'con tecnologías web estándar (HTML, CSS y JavaScript) y utiliza Supabase como base de datos '
    'y sistema de almacenamiento en la nube.')

add_heading(doc, '¿Qué hace el sitio?', 2)
bullets_resumen = [
    'Muestra un carrusel de imágenes en la sección Hero con texto dinámico.',
    'Lista los horarios de servicios semanales (obtenidos en tiempo real de la base de datos).',
    'Presenta a los pastores generales con su foto actualizable.',
    'Publica anuncios de la iglesia con imagen, fecha y descripción.',
    'Muestra una galería de fotos de la comunidad.',
    'Permite registrarse al Programa Alpha (curso de introducción a la fe cristiana).',
    'Permite enviar peticiones de oración con datos de contacto opcionales.',
    'Muestra la ubicación con un mapa embebido de Google Maps.',
    'Incluye una sección de donaciones con enlace configurable.',
]
for b in bullets_resumen:
    add_bullet(doc, b)

add_heading(doc, 'Arquitectura del sistema', 2)
add_body(doc,
    'El proyecto sigue una arquitectura de tres capas bien separadas:')

add_heading(doc, 'Capa 1 — Frontend (lo que ve el usuario)', 3)
add_body(doc,
    'Archivos HTML, CSS y JavaScript puros que se sirven de forma estática. '
    'No requieren un servidor propio; pueden hospedarse en GitHub Pages, Netlify, '
    'Vercel o cualquier hosting estático de forma gratuita.')

add_heading(doc, 'Capa 2 — Base de datos y almacenamiento (Supabase)', 3)
add_body(doc,
    'Supabase es un servicio en la nube que proporciona:')
add_bullet(doc, 'Base de datos PostgreSQL (tablas de horarios, anuncios, galería, etc.).')
add_bullet(doc, 'Storage (almacenamiento de imágenes en buckets).')
add_bullet(doc, 'Autenticación (para proteger el panel de administración).')
add_bullet(doc, 'API REST automática (el JavaScript llama a Supabase directamente, sin servidor intermedio).')

add_heading(doc, 'Capa 3 — Panel de Administración', 3)
add_body(doc,
    'Los archivos admin.html y admin.js conforman un panel privado donde el equipo de la iglesia '
    'puede gestionar todo el contenido del sitio sin tocar código: publicar anuncios, '
    'subir fotos, administrar horarios, ver registros Alpha y peticiones de oración.')

add_heading(doc, 'Flujo de datos', 2)
add_code(doc,
"""  [Visitante abre el sitio]
         |
         v
  index.html se carga en el navegador
         |
         v
  index.js llama a Supabase (API REST sobre HTTPS)
         |
         v
  Supabase devuelve datos (horarios, anuncios, fotos...)
         |
         v
  JavaScript actualiza el HTML dinámicamente (DOM)
         |
         v
  El visitante ve el contenido actualizado""")

add_heading(doc, 'Tecnologías utilizadas', 2)
make_table(doc,
    ['Tecnología', 'Versión / CDN', 'Para qué sirve'],
    [
        ['HTML5', 'Estándar', 'Estructura y contenido de las páginas'],
        ['CSS3', 'Estándar', 'Estilos, colores, tipografía y diseño responsive'],
        ['JavaScript (ES2022+)', 'Nativo del navegador', 'Lógica dinámica, llamadas a Supabase, formularios'],
        ['Supabase JS', 'CDN v2 (jsDelivr)', 'Cliente para conectar con la base de datos Supabase'],
        ['Google Fonts', 'CDN', 'Fuentes Bebas Neue, DM Sans, Dancing Script'],
        ['GitHub Pages', '—', 'Hosting estático gratuito del sitio público'],
        ['Supabase', 'Cloud (plan gratuito)', 'Base de datos, storage y autenticación'],
    ],
    col_widths=[4.5, 4, 7.5])

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 3. ARCHIVOS DEL PROYECTO
# ═══════════════════════════════════════════════════════

add_heading(doc, '2. Archivos del Proyecto', 1)
add_body(doc,
    'La siguiente tabla describe todos los archivos que conforman el proyecto y la función '
    'de cada uno. Entender qué contiene cada archivo es el primer paso para poder modificar el sitio.')

make_table(doc,
    ['Archivo', 'Tipo', 'Para qué sirve'],
    [
        ['index.html', 'HTML', 'Página principal pública. Contiene toda la estructura visual: nav, hero, secciones, formularios y footer.'],
        ['style.css', 'CSS', 'Hoja de estilos global. Define colores, tipografías, layout y diseño responsive de index.html.'],
        ['index.js', 'JavaScript', 'Lógica del sitio público: carga datos de Supabase y maneja los formularios Alpha y oración.'],
        ['admin.html', 'HTML', 'Panel de administración privado. Contiene la interfaz del panel con sidebar y todos los tabs.'],
        ['admin.js', 'JavaScript', 'Lógica del panel admin: autenticación, CRUD de contenido, subida de imágenes, exportación CSV.'],
        ['logo.png', 'Imagen', 'Logo de la iglesia en versión oscura (se usa en la barra de navegación).'],
        ['logo_inv.png', 'Imagen', 'Logo en versión clara/invertida (se usa en el hero y el footer sobre fondos oscuros).'],
        ['1.png', 'Imagen', 'Icono de la pestaña del navegador (favicon).'],
        ['pastores.jpg', 'Imagen', 'Foto de respaldo de los pastores generales (puede actualizarse desde el panel admin).'],
    ],
    col_widths=[3.5, 2.5, 10])

add_note(doc,
    'Los archivos de imagen (logo.png, pastores.jpg, etc.) son el punto de partida. '
    'Las fotos dinámicas (galería, hero, pastor) se gestionan desde el panel admin '
    'y se almacenan en Supabase Storage, no en el repositorio.', 'nota')

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 4. ESTRUCTURA HTML
# ═══════════════════════════════════════════════════════

add_heading(doc, '3. Estructura HTML (index.html)', 1)
add_body(doc,
    'El archivo index.html está organizado en secciones semánticas. Cada sección tiene un '
    'identificador (id="...") que permite navegar directamente a ella desde los enlaces del menú '
    'o desde los botones del hero.')

# NAV
add_heading(doc, '3.1 Barra de Navegación (<nav>)', 2)
add_body(doc,
    'La barra de navegación es fija en la parte superior de la pantalla. Contiene:')
add_bullet(doc, 'Logo de la iglesia (logo.png) y nombre "Linaje Escogido" con subtítulo "Iglesia".')
add_bullet(doc, 'Lista de enlaces (<ul class="nav-links">) que lleva a cada sección de la página.')
add_bullet(doc, 'Botón "Visítanos" con estilo primario (btn-primary).')
add_code(doc,
"""<nav>
  <div class="nav-brand">
    <img src="logo.png" alt="Linaje Escogido" class="nav-logo"/>
    <div>
      <div class="nav-name">Linaje Escogido</div>
      <div class="nav-sub">Iglesia</div>
    </div>
  </div>
  <ul class="nav-links">
    <li><a href="#horarios">Horarios</a></li>
    <li><a href="#pastores">Pastores</a></li>
    <!-- ... más enlaces ... -->
  </ul>
  <button class="btn-primary">Visítanos</button>
</nav>""")
add_body(doc, 'Para agregar un nuevo enlace al menú, copia uno de los elementos <li> y '
    'cambia el href al id de la nueva sección y el texto visible.')
add_note(doc, 'Si cambias el id de una sección en el HTML, recuerda actualizar el enlace correspondiente en el <nav> y en los botones del hero.', 'aviso')

# HERO
add_heading(doc, '3.2 Sección Hero (<section class="hero">)', 2)
add_body(doc,
    'El hero es la primera sección grande que ve el visitante. Ocupa toda la pantalla y contiene:')
add_bullet(doc, 'Logo invertido (logo_inv.png) de gran tamaño.')
add_bullet(doc, 'Título principal: "Linaje Escogido".')
add_bullet(doc, 'Dos párrafos con clase hero-tagline: el subtítulo de la iglesia y la ciudad.')
add_bullet(doc, 'Dos botones: "Cómo llegar" (lleva a #ubicacion) y "Programa Alpha →" (lleva a #alpha).')
add_body(doc,
    'El fondo del hero es un carrusel de imágenes que se carga dinámicamente desde Supabase '
    '(tabla hero_fotos). Si no hay fotos en la base de datos, el fondo queda oscuro por el CSS.')
add_code(doc,
"""<section class="hero">
  <img src="logo_inv.png" alt="Linaje Escogido" class="hero-logo"/>
  <h1>Linaje Escogido</h1>
  <p class="hero-tagline">Una Iglesia Funcional · Apasionada · Sensible al Espíritu Santo</p>
  <p class="hero-tagline hero-tagline-place">Reynosa, Tamaulipas</p>
  <div class="hero-btns">
    <button class="btn-cream"
      onclick="document.getElementById('ubicacion').scrollIntoView({behavior:'smooth'})">
      Cómo llegar
    </button>
    <button class="btn-outline-light"
      onclick="document.getElementById('alpha').scrollIntoView({behavior:'smooth'})">
      Programa Alpha →
    </button>
  </div>
</section>""")
add_note(doc,
    'El texto "Una Iglesia Funcional · Apasionada · Sensible al Espíritu Santo" y la ciudad '
    '"Reynosa, Tamaulipas" se pueden actualizar desde el Panel Admin en la pestaña Hero, '
    'sin necesidad de editar el HTML.', 'tip')

# HORARIOS
add_heading(doc, '3.3 Sección Horarios (<section id="horarios">)', 2)
add_body(doc,
    'Esta sección muestra los servicios semanales con día, hora y nombre del servicio. '
    'El contenido es completamente dinámico: se carga desde la tabla horarios de Supabase.')
add_code(doc,
"""<section id="horarios">
  <div class="label">Servicios</div>
  <div class="title">Horarios</div>
  <div class="horarios-grid" id="horariosGrid">
    <!-- JavaScript llena este div con las tarjetas de horario -->
    <div class="horario-card" style="opacity:.4;font-size:13px;">Cargando...</div>
  </div>
</section>""")
add_body(doc,
    'El elemento con id="horariosGrid" es el contenedor donde el JavaScript '
    'inyecta las tarjetas de horario. El div "Cargando..." es solo un placeholder '
    'temporal que desaparece cuando llegan los datos.')
add_note(doc, 'Para agregar, editar o eliminar horarios, ve al Panel Admin → pestaña Horarios. No es necesario editar el HTML.', 'tip')

# PASTORES
add_heading(doc, '3.4 Sección Pastores (<section id="pastores">)', 2)
add_body(doc,
    'La sección de pastores tiene dos partes: una foto de fondo a pantalla completa con overlay '
    'oscuro y una tarjeta de contenido con los nombres de los pastores generales.')
add_bullet(doc, 'La imagen de fondo se carga dinámicamente desde la tabla pastores_foto.')
add_bullet(doc, 'Si la imagen falla, se usa una imagen placeholder de placehold.co.')
add_bullet(doc, 'La frase en cursiva "¡Tenemos un lugar para ti!" usa la clase pastor-script (fuente Dancing Script).')
add_bullet(doc, 'Cada pastor tiene un pastor-card con su nombre y rol.')
add_code(doc,
"""<div class="pastor-grid">
  <div class="pastor-card">
    <div class="pastor-card-name">Pr. Hebert A. Gaytan Olvera</div>
    <div class="pastor-card-role">Pastor General</div>
  </div>
  <div class="pastor-card">
    <div class="pastor-card-name">Pra. Luisa Rangel Segovia</div>
    <div class="pastor-card-role">Pastora General</div>
  </div>
</div>""")
add_body(doc, 'Para cambiar los nombres de los pastores, edita directamente los textos '
    'dentro de los div.pastor-card-name en el HTML.')

# ANUNCIOS
add_heading(doc, '3.5 Sección Anuncios (<section id="anuncios">)', 2)
add_body(doc,
    'Muestra los 3 anuncios más recientes activos, con imagen, fecha, título y descripción. '
    'El contenedor es #anunciosGrid y es llenado por la función loadAnuncios() en index.js.')
add_code(doc,
"""<section id="anuncios">
  <div class="label">Novedades</div>
  <div class="title">Anuncios</div>
  <div class="anuncios-grid" id="anunciosGrid">
    <!-- Llenado dinámicamente por index.js -->
  </div>
</section>""")
add_note(doc, 'La sección de Anuncios siempre mostrará máximo 3 anuncios (los más recientes). Para publicar o despublicar anuncios, usa el Panel Admin.', 'nota')

# GALERÍA
add_heading(doc, '3.6 Sección Galería', 2)
add_body(doc,
    'La galería de fotos de la comunidad se carga desde la tabla galeria_fotos. '
    'Las fotos se muestran en un grid responsivo. Al final hay un enlace al Instagram oficial.')
add_code(doc,
"""<div class="gallery-grid" id="galeriaGrid">
  <!-- Imágenes inyectadas por loadGaleria() -->
</div>
<div style="text-align:center;margin-top:28px;">
  <a href="https://instagram.com/linajeescogido.mx" target="_blank" class="gallery-link">
    @linajeescogido.mx en Instagram →
  </a>
</div>""")
add_body(doc, 'Para cambiar el usuario de Instagram, busca el texto "@linajeescogido.mx" '
    'en el HTML y actualiza la URL del href y el texto visible.')

# ALPHA + ORACIÓN
add_heading(doc, '3.7 Sección Formularios: Alpha y Oración (<section id="alpha">)', 2)
add_body(doc,
    'Esta sección contiene dos formularios lado a lado (en pantallas grandes):')

add_heading(doc, 'Formulario Alpha', 3)
add_body(doc, 'Recopila: Nombre, Apellidos, Correo electrónico y Teléfono. '
    'Al enviarlo, se guarda en la tabla registros_alpha de Supabase. '
    'El formulario llama a la función handleAlpha(event) al hacer submit.')
add_bullet(doc, 'id="alphaNombre" — Campo de nombre')
add_bullet(doc, 'id="alphaApellidos" — Campo de apellidos')
add_bullet(doc, 'id="alphaEmail" — Campo de correo')
add_bullet(doc, 'id="alphaTelefono" — Campo de teléfono (opcional)')

add_heading(doc, 'Formulario de Petición de Oración', 3)
add_body(doc, 'Recopila información más detallada:')
add_bullet(doc, 'Texto de la petición (textarea, obligatorio)')
add_bullet(doc, 'Nombre y Apellidos')
add_bullet(doc, 'Email y Teléfono (opcionales)')
add_bullet(doc, '¿De qué parte del mundo nos ves? (ciudad/país)')
add_bullet(doc, '¿Cómo te enteraste?')
add_bullet(doc, '¿Tomaste hoy la decisión de seguir a Jesús? (select)')
add_bullet(doc, 'Checkbox de consentimiento de datos (obligatorio)')
add_body(doc, 'Al enviarlo, se guarda en la tabla peticiones_oracion de Supabase.')
add_note(doc,
    'Ambos formularios tienen validación HTML nativa (atributo required). '
    'Si un campo requerido está vacío, el navegador muestra un aviso automáticamente.', 'nota')

# UBICACIÓN
add_heading(doc, '3.8 Sección Ubicación (<section id="ubicacion">)', 2)
add_body(doc,
    'Contiene un iframe de Google Maps embebido y la dirección física de la iglesia.')
add_code(doc,
"""<section id="ubicacion">
  <div class="label">Encuéntranos</div>
  <div class="title">Ubicación</div>
  <div class="map-embed">
    <iframe src="https://www.google.com/maps/embed?..."
            width="600" height="450"
            style="border:0;" allowfullscreen="" loading="lazy">
    </iframe>
  </div>
  <div class="map-address">
    <strong>Calle Ávila Camacho #130</strong><br>
    Col. La Presa, Reynosa, Tamaulipas 88750
  </div>
</section>""")
add_body(doc, 'Para actualizar la dirección en el mapa:')
add_numbered(doc, 'Ve a maps.google.com y busca la nueva dirección.')
add_numbered(doc, 'Haz clic en "Compartir" → "Insertar un mapa" y copia el código del iframe.')
add_numbered(doc, 'Reemplaza el atributo src del <iframe> en el HTML.')
add_numbered(doc, 'Actualiza también el texto en div.map-address si cambió la dirección.')

# DONACIONES
add_heading(doc, '3.9 Sección Donaciones (<section class="donaciones-section">)', 2)
add_body(doc,
    'Sección oscura con texto motivacional y un botón "Donar ahora". '
    'El enlace del botón (href="#") debe actualizarse con la URL de la plataforma de donaciones '
    'que use la iglesia (por ejemplo, un enlace de PayPal, Mercado Pago, etc.).')
add_code(doc,
"""<a href="#" class="btn-cream" ...>Donar ahora</a>
<!-- Reemplaza el "#" con la URL real de donaciones -->""")

# FOOTER
add_heading(doc, '3.10 Footer (<footer>)', 2)
add_body(doc,
    'El pie de página contiene el logo invertido, el copyright y los enlaces a redes sociales. '
    'Muestra: © 2026 Linaje Escogido · Reynosa, Tamaulipas.')
add_code(doc,
"""<footer>
  <div class="footer-brand">
    <img src="logo_inv.png" alt="Linaje Escogido" class="footer-logo"/>
    <span class="footer-copy">© 2026 Linaje Escogido · Reynosa, Tamaulipas</span>
  </div>
  <div class="footer-social">
    <a href="https://www.facebook.com/..." target="_blank">Facebook</a>
    <a href="https://instagram.com/linajeescogido.mx" target="_blank">Instagram</a>
  </div>
</footer>""")
add_body(doc, 'Para actualizar los enlaces de redes sociales, cambia los valores del atributo href.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 5. ESTILOS CSS
# ═══════════════════════════════════════════════════════

add_heading(doc, '4. Estilos CSS (style.css)', 1)
add_body(doc,
    'El archivo style.css controla toda la apariencia visual del sitio: colores, tipografía, '
    'espaciado, layout y adaptación a distintos tamaños de pantalla (responsive design).')

add_heading(doc, '4.1 Variables CSS (Custom Properties)', 2)
add_body(doc,
    'Las variables CSS son el sistema de colores y valores reutilizables del sitio. '
    'Están definidas en :root (raíz del documento) para que estén disponibles globalmente. '
    'Cambiar una variable actualiza automáticamente todos los elementos que la usan.')
add_code(doc,
""":root {
  --cream:  #F5EFE6;              /* Color crema — fondo principal del sitio */
  --dark:   #1A1A1A;              /* Casi negro — texto y fondos oscuros      */
  --mid:    #888;                 /* Gris medio — textos secundarios          */
  --border: rgba(26,26,26,0.15); /* Borde sutil — separadores y contornos    */
}""")
add_body(doc, 'Para usar una variable en cualquier regla CSS:')
add_code(doc,
""".mi-elemento {
  color: var(--dark);
  background: var(--cream);
  border: 1px solid var(--border);
}""")

add_heading(doc, 'Tabla de variables de color', 3)
make_table(doc,
    ['Variable', 'Valor hex', 'Dónde se usa'],
    [
        ['--cream', '#F5EFE6', 'Fondo general del sitio, botones claros, texto sobre fondos oscuros'],
        ['--dark', '#1A1A1A', 'Texto principal, fondo de nav, sección pastores, donaciones, footer'],
        ['--mid', '#888', 'Textos secundarios, metadatos, fechas de anuncios'],
        ['--border', 'rgba(26,26,26,0.15)', 'Bordes de tarjetas, separadores, contornos de inputs'],
    ],
    col_widths=[3.5, 3, 9.5])

add_heading(doc, '4.2 Tipografías', 2)
add_body(doc,
    'El sitio usa tres familias tipográficas de Google Fonts, cargadas en el <head> del HTML:')
make_table(doc,
    ['Fuente', 'Estilo', 'Dónde se usa'],
    [
        ['Bebas Neue', 'Display, solo mayúsculas', 'Títulos grandes de secciones (.title), título del hero'],
        ['DM Sans', 'Sans-serif, múltiples pesos', 'Todo el texto de cuerpo, navegación, formularios, botones'],
        ['Dancing Script', 'Manuscrita cursiva', 'Frase destacada en la sección de Pastores (.pastor-script)'],
    ],
    col_widths=[4, 4, 8])
add_body(doc, 'Para cambiar una fuente, reemplaza el nombre en el link de Google Fonts del <head> '
    'y actualiza la propiedad font-family correspondiente en el CSS.')

add_heading(doc, '4.3 Estructura de estilos por sección', 2)

sections_css = [
    ('Navegación (nav, .nav-brand, .nav-links)',
     'Define la barra superior fija (position: fixed), con fondo --dark, '
     'flexbox para alinear logo, enlaces y botón. En mobile se ocultan los enlaces '
     'y se reorganizan los elementos. La clase .nav-logo controla el tamaño del logo.'),
    ('Hero (.hero, .hero-logo, .hero-tagline, .hero-btns)',
     'El hero usa position: relative con height: 100vh para ocupar toda la pantalla. '
     'Las fotos de fondo se posicionan absolutas con z-index bajo. '
     'El contenido (logo, texto, botones) tiene z-index alto para superponerse. '
     'Un overlay oscuro (::before o elemento .hero-overlay) asegura legibilidad del texto.'),
    ('Horarios (.horarios-grid, .horario-card)',
     'Grid CSS de 3 columnas en desktop, 1 columna en mobile. '
     'Las tarjetas alternan entre estilo claro (fondo crema) y oscuro (fondo dark), '
     'generado desde JavaScript al crear las cards.'),
    ('Pastores (.pastores-section, .pastor-img-wrap, .pastor-content)',
     'Layout de dos columnas: imagen a pantalla completa a la izquierda, '
     'contenido textual a la derecha. En mobile apila verticalmente. '
     'La foto tiene un overlay de gradiente para mejorar la legibilidad del texto superpuesto.'),
    ('Anuncios (.anuncios-grid, .anuncio)',
     'Grid de hasta 3 columnas. Cada .anuncio es una tarjeta con imagen, '
     'fecha en pequeño, título y descripción. Borde sutil con --border.'),
    ('Galería (.gallery-grid, .gallery-item)',
     'Grid CSS con columnas automáticas usando repeat(auto-fill, minmax(200px, 1fr)). '
     'Las imágenes tienen object-fit: cover para que no se distorsionen.'),
    ('Formularios (.form-block, .field-wrap, .field-input, .field-textarea)',
     'Los dos formularios usan display: flex con flex-wrap para adaptarse al ancho. '
     'Inputs y textareas tienen estilos unificados con --border y radio de borde. '
     'Las etiquetas (.field-label) están en pequeño y en gris.'),
    ('Botones (.btn-primary, .btn-cream, .btn-outline-light, .btn-submit-dark)',
     'Cada variante de botón tiene su propio color de fondo, texto y hover. '
     '--btn-primary: fondo crema sobre oscuro. '
     '--btn-cream: fondo #F5EFE6. '
     '--btn-outline-light: solo borde claro, fondo transparente. '
     '--btn-submit-dark: fondo oscuro para formularios.'),
    ('Donaciones (.donaciones-section, .don-title)',
     'Sección de fondo oscuro con texto centrado y color crema. '
     'Usa padding generoso y un tamaño de letra grande para el título.'),
    ('Footer (footer, .footer-brand, .footer-social)',
     'Fondo oscuro, layout flexbox en dos columnas: marca/copyright a la izquierda, '
     'redes sociales a la derecha. En mobile se apila verticalmente y centra.'),
]

for title_css, desc_css in sections_css:
    add_heading(doc, title_css, 3)
    add_body(doc, desc_css)

add_heading(doc, '4.4 Diseño Responsive (Media Queries)', 2)
add_body(doc,
    'El sitio se adapta a diferentes tamaños de pantalla usando @media queries. '
    'El enfoque es "desktop-first": los estilos base son para pantallas grandes '
    'y las media queries ajustan para pantallas pequeñas.')
add_code(doc,
"""/* Ejemplo de media query para tablets y móviles */
@media (max-width: 768px) {
  .nav-links {
    display: none;           /* Oculta el menú en mobile */
  }
  .horarios-grid {
    grid-template-columns: 1fr; /* 1 columna en mobile */
  }
  .pastores-section {
    flex-direction: column;  /* Apila en lugar de lado a lado */
  }
}""")
add_body(doc, 'Los breakpoints principales utilizados son:')
make_table(doc,
    ['Breakpoint', 'Ancho máximo', 'Qué cambia'],
    [
        ['Tablet', '1024px', 'Se ajustan grids a 2 columnas, se reduce tamaño de fuente en hero'],
        ['Mobile', '768px', 'Se oculta nav-links, layouts pasan a 1 columna, formularios se apilan'],
        ['Mobile pequeño', '480px', 'Se reducen paddings y márgenes, fuentes se ajustan con clamp()'],
    ],
    col_widths=[3.5, 3.5, 9])

add_heading(doc, '4.5 La función clamp() para tipografía fluida', 2)
add_body(doc,
    'El sitio usa clamp() para que los títulos escalen fluidamente entre el tamaño mínimo '
    'y máximo según el ancho de la ventana, sin necesidad de media queries específicas.')
add_code(doc,
"""/* Sintaxis: clamp(mínimo, preferido, máximo) */
.title {
  font-size: clamp(32px, 5vw, 56px);
  /* En mobile será 32px, en desktop 56px,
     y en medio escala proporcionalmente */
}""")
add_note(doc,
    'vw significa "viewport width" — 5vw es el 5% del ancho de la ventana. '
    'Si la ventana tiene 1000px de ancho, 5vw = 50px.', 'nota')

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 6. LÓGICA JAVASCRIPT
# ═══════════════════════════════════════════════════════

add_heading(doc, '5. Lógica JavaScript (index.js)', 1)
add_body(doc,
    'El archivo index.js es el cerebro dinámico del sitio. Se encarga de conectar con Supabase '
    'para cargar contenido y de procesar los formularios del visitante.')

add_heading(doc, '5.1 Configuración de Supabase', 2)
add_body(doc,
    'Las primeras líneas del archivo configuran la conexión con Supabase:')
add_code(doc,
"""const SUPABASE_URL = 'https://xzwyyzgcxhpbjioqdrmn.supabase.co';
const SUPABASE_KEY = 'eyJ...'; // Clave pública anon (safe para el frontend)

const { createClient } = supabase; // supabase viene del CDN en el HTML
const db = createClient(SUPABASE_URL, SUPABASE_KEY);""")
add_note(doc,
    'La SUPABASE_KEY que aparece en el frontend es la clave "anon" (anónima). '
    'Es pública por diseño — no es un secreto. La seguridad real se configura '
    'con Row Level Security (RLS) directamente en Supabase, que define qué puede '
    'leer/escribir un usuario no autenticado.', 'nota')

add_heading(doc, '5.2 Inicialización (DOMContentLoaded)', 2)
add_body(doc,
    'Cuando el HTML termina de cargarse en el navegador, se dispara el evento '
    'DOMContentLoaded y se llaman todas las funciones de carga:')
add_code(doc,
"""document.addEventListener('DOMContentLoaded', () => {
  loadAnuncios();    // Carga los anuncios de la iglesia
  loadGaleria();     // Carga las fotos de la galería
  loadHero();        // Carga las fotos del carrusel hero
  loadHeroTexto();   // Carga el subtítulo y ciudad del hero
  loadPastorFoto();  // Carga la foto de los pastores
  loadHorarios();    // Carga los horarios de servicios
});""")
add_body(doc, 'Todas las funciones de carga son asíncronas (async/await), lo que significa '
    'que el resto de la página no espera a que terminen — todo carga en paralelo.')

add_heading(doc, '5.3 Función loadHero()', 2)
add_body(doc, 'Carga las fotos del carrusel de la sección hero.')
make_table(doc,
    ['Aspecto', 'Detalle'],
    [
        ['Tabla Supabase', 'hero_fotos'],
        ['Filtro', 'activa = true'],
        ['Orden', 'Por columna "orden" ascendente'],
        ['Qué hace', 'Crea elementos <img> posicionados absolutamente. Si hay más de 1 foto, rota cada 5 segundos con animación de opacidad (fade).'],
        ['Cuándo se ejecuta', 'Al cargar la página (DOMContentLoaded)'],
    ],
    col_widths=[4, 12])
add_code(doc,
"""async function loadHero() {
  const { data: fotos } = await db
    .from('hero_fotos')
    .select('*')
    .eq('activa', true)
    .order('orden');

  if (!fotos || fotos.length === 0) return;

  // Crea y agrega las imágenes al DOM
  fotos.forEach((foto, i) => {
    const img = document.createElement('img');
    img.src = foto.url;
    img.className = 'hero-bg-img';
    img.style.opacity = i === 0 ? '1' : '0'; // Solo la primera visible
    heroSection.appendChild(img);
  });

  // Si hay más de una, inicia el carrusel
  if (fotos.length > 1) {
    let current = 0;
    setInterval(() => {
      imgs[current].style.opacity = '0';
      current = (current + 1) % imgs.length;
      imgs[current].style.opacity = '1';
    }, 5000); // Cambia cada 5 segundos
  }
}""")

add_heading(doc, '5.4 Función loadHeroTexto()', 2)
add_body(doc, 'Carga el subtítulo y la ciudad de la iglesia desde la base de datos.')
make_table(doc,
    ['Aspecto', 'Detalle'],
    [
        ['Tabla Supabase', 'hero_contenido'],
        ['Qué obtiene', 'Los campos "subtitulo" y "ubicacion"'],
        ['Qué hace', 'Busca los elementos .hero-tagline en el DOM y reemplaza su texto con los valores de la BD'],
        ['Para qué sirve', 'Permite cambiar el lema de la iglesia o la ciudad sin editar el código HTML'],
    ],
    col_widths=[4, 12])

add_heading(doc, '5.5 Función loadPastorFoto()', 2)
add_body(doc, 'Actualiza la foto de los pastores desde la base de datos.')
make_table(doc,
    ['Aspecto', 'Detalle'],
    [
        ['Tabla Supabase', 'pastores_foto'],
        ['Filtro', 'activa = true'],
        ['Orden', 'Por created_at descendente (la más reciente primero)'],
        ['Qué hace', 'Toma la URL de la primera foto activa y la asigna como src del elemento img.pastor-img'],
    ],
    col_widths=[4, 12])

add_heading(doc, '5.6 Función loadHorarios()', 2)
add_body(doc, 'Carga y renderiza las tarjetas de horarios de servicios.')
make_table(doc,
    ['Aspecto', 'Detalle'],
    [
        ['Tabla Supabase', 'horarios'],
        ['Filtro', 'activo = true'],
        ['Orden', 'Por columna "orden" ascendente'],
        ['Qué hace', 'Genera HTML dinámico para cada horario. Las tarjetas pares tienen fondo claro, las impares fondo oscuro.'],
    ],
    col_widths=[4, 12])
add_code(doc,
"""async function loadHorarios() {
  const { data: horarios } = await db
    .from('horarios')
    .select('*')
    .eq('activo', true)
    .order('orden');

  const grid = document.getElementById('horariosGrid');
  grid.innerHTML = ''; // Limpia el placeholder

  horarios.forEach((h, i) => {
    const card = document.createElement('div');
    // Alterna entre estilo claro y oscuro
    card.className = i % 2 === 0 ? 'horario-card' : 'horario-card dark';
    card.innerHTML = `
      <div class="horario-titulo">${h.titulo}</div>
      <div class="horario-dia">${h.dia}</div>
      <div class="horario-hora">${h.hora}</div>
    `;
    grid.appendChild(card);
  });
}""")

add_heading(doc, '5.7 Función loadAnuncios()', 2)
add_body(doc, 'Carga los 3 anuncios más recientes activos.')
make_table(doc,
    ['Aspecto', 'Detalle'],
    [
        ['Tabla Supabase', 'anuncios'],
        ['Filtro', 'activo = true'],
        ['Orden', 'Por fecha descendente (más reciente primero)'],
        ['Límite', '3 anuncios'],
        ['Qué hace', 'Genera cards con imagen (si tiene), fecha formateada, título y cuerpo del anuncio'],
    ],
    col_widths=[4, 12])

add_heading(doc, '5.8 Función loadGaleria()', 2)
add_body(doc, 'Llena la galería de fotos de la comunidad.')
make_table(doc,
    ['Aspecto', 'Detalle'],
    [
        ['Tabla Supabase', 'galeria_fotos'],
        ['Filtro', 'activa = true'],
        ['Orden', 'Por columna "orden" ascendente'],
        ['Qué hace', 'Crea elementos <img> dentro de <div class="gallery-item"> y los agrega al #galeriaGrid'],
    ],
    col_widths=[4, 12])

add_heading(doc, '5.9 Función handleAlpha(e)', 2)
add_body(doc, 'Procesa el envío del formulario de registro Alpha.')
add_code(doc,
"""async function handleAlpha(e) {
  e.preventDefault(); // Previene la recarga de página

  const nombre    = document.getElementById('alphaNombre').value;
  const apellidos = document.getElementById('alphaApellidos').value;
  const email     = document.getElementById('alphaEmail').value;
  const telefono  = document.getElementById('alphaTelefono').value;

  const { error } = await db
    .from('registros_alpha')
    .insert([{ nombre, apellidos, email, telefono }]);

  if (error) {
    alert('Hubo un error. Inténtalo de nuevo.');
  } else {
    alert('¡Registro exitoso! Te contactaremos pronto.');
    e.target.reset(); // Limpia el formulario
  }
}""")

add_heading(doc, '5.10 Función handleOracion(e)', 2)
add_body(doc, 'Procesa el envío del formulario de petición de oración.')
make_table(doc,
    ['Campo guardado', 'Descripción'],
    [
        ['nombre', 'Nombre del solicitante'],
        ['apellidos', 'Apellidos (opcional)'],
        ['email', 'Correo electrónico (opcional)'],
        ['telefono', 'Teléfono (opcional)'],
        ['peticion', 'Texto de la petición de oración (obligatorio)'],
        ['ubicacion', '¿De qué parte del mundo nos ves?'],
        ['referencia', '¿Cómo te enteraste?'],
        ['decision_seguimiento', 'Respuesta al select de decisión de seguir a Jesús'],
    ],
    col_widths=[5, 11])

add_heading(doc, '5.11 Función formatDate(dateStr)', 2)
add_body(doc,
    'Función utilitaria que convierte una fecha en formato ISO (YYYY-MM-DD) '
    'a un formato legible en español.')
add_code(doc,
"""function formatDate(dateStr) {
  // Entrada:  '2026-05-15'
  // Salida:   '15 de mayo de 2026'
  const [year, month, day] = dateStr.split('-');
  const meses = [
    'enero','febrero','marzo','abril','mayo','junio',
    'julio','agosto','septiembre','octubre','noviembre','diciembre'
  ];
  return `${parseInt(day)} de ${meses[parseInt(month)-1]} de ${year}`;
}""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 7. PANEL DE ADMINISTRACIÓN
# ═══════════════════════════════════════════════════════

add_heading(doc, '6. Panel de Administración (admin.html + admin.js)', 1)
add_body(doc,
    'El panel de administración es una interfaz privada desde la cual el equipo de la iglesia '
    'puede gestionar todo el contenido del sitio sin necesidad de escribir código. '
    'Solo pueden acceder quienes tengan una cuenta de usuario creada en Supabase Auth.')

add_heading(doc, '6.1 Cómo acceder', 2)
add_numbered(doc, 'Abre el archivo admin.html en el navegador (o navega a la URL del admin).')
add_numbered(doc, 'Ingresa tu correo electrónico y contraseña (configurados en Supabase Auth).')
add_numbered(doc, 'Haz clic en "Iniciar sesión".')
add_numbered(doc, 'Si las credenciales son correctas, se mostrará el panel completo.')
add_note(doc,
    'Para crear nuevos usuarios del admin, ve al panel de Supabase → Authentication → Users → Invite User. '
    'Esto envía un correo de invitación a la persona.', 'tip')

add_heading(doc, '6.2 Estructura del panel', 2)
add_body(doc,
    'El panel tiene dos zonas principales:')
add_bullet(doc, 'Sidebar (barra lateral): lista de pestañas para navegar entre secciones.')
add_bullet(doc, 'Área de contenido (main): zona donde se muestra y gestiona el contenido de cada pestaña.')
add_body(doc,
    'En dispositivos móviles, el sidebar se convierte en un menú drawer que se abre '
    'con un botón de hamburguesa (≡) en la esquina superior.')

add_heading(doc, '6.3 Pestaña: Anuncios', 2)
add_body(doc, 'Permite gestionar los anuncios que aparecen en el sitio público.')
make_table(doc,
    ['Función', 'Qué hace'],
    [
        ['loadAnuncios()', 'Carga todos los anuncios en una tabla del panel, mostrando título, fecha, estado activo/inactivo y acciones.'],
        ['handleAnuncio(e)', 'Crea un nuevo anuncio. Toma: título, fecha, descripción e imagen opcional. Si hay imagen, la sube primero al bucket "fotos" de Supabase Storage y guarda la URL pública.'],
        ['toggleAnuncio(id, activo)', 'Activa o desactiva un anuncio (sin borrarlo). Los anuncios inactivos no se muestran en el sitio público.'],
        ['deleteAnuncio(id, imagenUrl)', 'Elimina el anuncio de la base de datos. Si tenía imagen, la elimina también del Storage.'],
        ['filterAnuncios()', 'Filtra la tabla de anuncios en tiempo real según el texto escrito en el campo de búsqueda.'],
        ['exportAlpha()', 'Exporta todos los registros Alpha a un archivo CSV descargable (disponible también desde esta pestaña por conveniencia).'],
    ],
    col_widths=[5, 11])

add_heading(doc, '6.4 Pestaña: Hero', 2)
add_body(doc, 'Gestiona las fotos del carrusel de la sección hero y el texto principal.')
make_table(doc,
    ['Función', 'Qué hace'],
    [
        ['loadHeroFotos()', 'Muestra un grid con todas las fotos del carrusel hero, indicando cuáles están activas y su orden.'],
        ['handleHeroUpload(e)', 'Sube una o varias fotos nuevas al bucket de Supabase y las registra en la tabla hero_fotos como activas.'],
        ['loadHeroTexto()', 'Carga el subtítulo y la ubicación actuales del hero en los campos de edición.'],
        ['handleHeroTexto(e)', 'Guarda los cambios del subtítulo y la ubicación en la tabla hero_contenido. Los cambios se reflejan en el sitio inmediatamente.'],
    ],
    col_widths=[5, 11])

add_heading(doc, '6.5 Pestaña: Galería', 2)
add_body(doc, 'Gestiona las fotos de la galería de la comunidad.')
make_table(doc,
    ['Función', 'Qué hace'],
    [
        ['loadGaleriaFotos()', 'Muestra un grid de todas las fotos de la galería con opciones de activar/desactivar y eliminar.'],
        ['handleGaleriaUpload(e)', 'Sube nuevas fotos a Supabase Storage y las registra en la tabla galeria_fotos.'],
        ['toggleFoto(tabla, id, activa)', 'Oculta o muestra una foto en el sitio público sin eliminarla. Funciona tanto para galería como para otras fotos.'],
        ['deleteFoto(tabla, id, url)', 'Elimina la foto de la base de datos y del Storage.'],
    ],
    col_widths=[5, 11])

add_heading(doc, '6.6 Pestaña: Horarios', 2)
add_body(doc, 'Gestiona los horarios de servicios de la iglesia.')
make_table(doc,
    ['Función', 'Qué hace'],
    [
        ['loadHorarios()', 'Muestra todos los horarios en una tabla con columnas: título, día, hora, estado y acciones.'],
        ['handleHorario(e)', 'Crea un nuevo horario con título (ej: "Servicio Dominical"), día (ej: "Domingo") y hora (ej: "10:00 AM").'],
        ['toggleHorario(id, activo)', 'Activa o desactiva un horario para el sitio público.'],
        ['deleteHorario(id)', 'Elimina permanentemente un horario.'],
    ],
    col_widths=[5, 11])

add_heading(doc, '6.7 Pestaña: Pastores', 2)
add_body(doc, 'Gestiona la foto principal de los pastores generales.')
make_table(doc,
    ['Función', 'Qué hace'],
    [
        ['loadPastorFoto()', 'Muestra la foto actual de los pastores y las fotos anteriores almacenadas.'],
        ['handlePastorUpload(e)', 'Sube una nueva foto de pastores. Si el archivo supera 2MB, lo comprime automáticamente antes de subirlo.'],
        ['deletePastorFoto(id, url)', 'Elimina una foto de pastores de la BD y del Storage.'],
        ['compressImage(file, maxMB)', 'Función interna que usa la Canvas API del navegador para redimensionar y comprimir imágenes grandes, preservando la relación de aspecto.'],
    ],
    col_widths=[5, 11])

add_heading(doc, '6.8 Pestaña: Registros Alpha', 2)
add_body(doc, 'Visualiza y exporta los registros del formulario Alpha del sitio público.')
make_table(doc,
    ['Función', 'Qué hace'],
    [
        ['loadAlphaRegistros()', 'Lista todos los registros en una tabla con: nombre, apellidos, email, teléfono y fecha de registro.'],
        ['exportAlpha()', 'Genera y descarga un archivo CSV con todos los registros para importar en Excel u hojas de cálculo.'],
    ],
    col_widths=[5, 11])

add_heading(doc, '6.9 Pestaña: Peticiones de Oración', 2)
add_body(doc, 'Visualiza y exporta las peticiones de oración enviadas desde el sitio.')
make_table(doc,
    ['Función', 'Qué hace'],
    [
        ['loadPeticiones()', 'Lista todas las peticiones con nombre, texto de la petición, ubicación, referencia y decisión de fe.'],
        ['exportPeticiones()', 'Descarga un CSV con todas las peticiones para revisión y seguimiento pastoral.'],
    ],
    col_widths=[5, 11])

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 8. BASE DE DATOS SUPABASE
# ═══════════════════════════════════════════════════════

add_heading(doc, '7. Base de Datos Supabase', 1)
add_body(doc,
    'Supabase usa PostgreSQL como motor de base de datos. Todas las tablas del proyecto '
    'están en el esquema "public" y tienen Row Level Security (RLS) habilitado para '
    'controlar quién puede leer o escribir datos.')

add_heading(doc, '7.1 Tabla: hero_fotos', 2)
add_body(doc, 'Almacena las fotos del carrusel de la sección hero.')
make_table(doc,
    ['Columna', 'Tipo', 'Descripción'],
    [
        ['id', 'uuid (PK)', 'Identificador único auto-generado'],
        ['url', 'text', 'URL pública de la imagen en Supabase Storage'],
        ['activa', 'boolean', 'Si es true, la foto aparece en el carrusel del sitio'],
        ['orden', 'integer', 'Número para ordenar las fotos en el carrusel (menor = primero)'],
        ['created_at', 'timestamptz', 'Fecha y hora de creación (auto-generado por Supabase)'],
    ],
    col_widths=[3.5, 3, 9.5])

add_heading(doc, '7.2 Tabla: hero_contenido', 2)
add_body(doc, 'Almacena el texto editable de la sección hero (subtítulo y ciudad).')
make_table(doc,
    ['Columna', 'Tipo', 'Descripción'],
    [
        ['id', 'uuid (PK)', 'Identificador único'],
        ['titulo', 'text', 'Título principal (actualmente no se usa dinámicamente)'],
        ['subtitulo', 'text', 'Frase debajo del título: "Una Iglesia Funcional · Apasionada..."'],
        ['ubicacion', 'text', 'Ciudad y estado: "Reynosa, Tamaulipas"'],
        ['created_at', 'timestamptz', 'Fecha de creación del registro'],
    ],
    col_widths=[3.5, 3, 9.5])
add_note(doc, 'Esta tabla típicamente tiene un solo registro que se actualiza. El código obtiene el primero disponible.', 'nota')

add_heading(doc, '7.3 Tabla: horarios', 2)
add_body(doc, 'Almacena los horarios de servicios de la iglesia.')
make_table(doc,
    ['Columna', 'Tipo', 'Descripción'],
    [
        ['id', 'uuid (PK)', 'Identificador único'],
        ['titulo', 'text', 'Nombre del servicio (ej: "Servicio Dominical", "Célula de Jóvenes")'],
        ['dia', 'text', 'Día de la semana (ej: "Domingo", "Miércoles")'],
        ['hora', 'text', 'Hora del servicio (ej: "10:00 AM", "7:30 PM")'],
        ['activo', 'boolean', 'Si es true, aparece en el sitio público'],
        ['orden', 'integer', 'Orden de aparición en la sección (menor = primero)'],
        ['created_at', 'timestamptz', 'Fecha de creación'],
    ],
    col_widths=[3.5, 3, 9.5])

add_heading(doc, '7.4 Tabla: anuncios', 2)
add_body(doc, 'Almacena los anuncios que aparecen en la sección "Novedades".')
make_table(doc,
    ['Columna', 'Tipo', 'Descripción'],
    [
        ['id', 'uuid (PK)', 'Identificador único'],
        ['titulo', 'text', 'Título del anuncio'],
        ['cuerpo', 'text', 'Descripción completa del anuncio'],
        ['fecha', 'date', 'Fecha del anuncio en formato YYYY-MM-DD'],
        ['imagen_url', 'text', 'URL de la imagen del anuncio (puede ser NULL si no tiene imagen)'],
        ['activo', 'boolean', 'Si es true, se muestra en el sitio (máx. 3 más recientes)'],
        ['created_at', 'timestamptz', 'Fecha de creación en la BD'],
    ],
    col_widths=[3.5, 3, 9.5])

add_heading(doc, '7.5 Tabla: galeria_fotos', 2)
add_body(doc, 'Almacena las fotos de la galería de la comunidad.')
make_table(doc,
    ['Columna', 'Tipo', 'Descripción'],
    [
        ['id', 'uuid (PK)', 'Identificador único'],
        ['url', 'text', 'URL pública de la foto en Storage'],
        ['activa', 'boolean', 'Si es true, la foto aparece en la galería del sitio'],
        ['orden', 'integer', 'Orden de aparición en la galería'],
        ['created_at', 'timestamptz', 'Fecha de creación'],
    ],
    col_widths=[3.5, 3, 9.5])

add_heading(doc, '7.6 Tabla: pastores_foto', 2)
add_body(doc, 'Almacena las fotos de los pastores generales.')
make_table(doc,
    ['Columna', 'Tipo', 'Descripción'],
    [
        ['id', 'uuid (PK)', 'Identificador único'],
        ['url', 'text', 'URL pública de la foto en Storage'],
        ['activa', 'boolean', 'Si es true, esta foto se muestra en el sitio'],
        ['orden', 'integer', 'Orden (el de menor número es el principal)'],
        ['created_at', 'timestamptz', 'Fecha de creación (se usa para obtener la más reciente)'],
    ],
    col_widths=[3.5, 3, 9.5])

add_heading(doc, '7.7 Tabla: registros_alpha', 2)
add_body(doc, 'Almacena los registros del formulario del Programa Alpha.')
make_table(doc,
    ['Columna', 'Tipo', 'Descripción'],
    [
        ['id', 'uuid (PK)', 'Identificador único'],
        ['nombre', 'text', 'Nombre del registrado'],
        ['apellidos', 'text', 'Apellidos del registrado'],
        ['email', 'text', 'Correo electrónico'],
        ['telefono', 'text', 'Número de teléfono (puede ser NULL)'],
        ['fecha_registro', 'timestamptz', 'Fecha y hora en que se registró (auto-generado)'],
    ],
    col_widths=[3.5, 3, 9.5])

add_heading(doc, '7.8 Tabla: peticiones_oracion', 2)
add_body(doc, 'Almacena las peticiones de oración enviadas desde el sitio.')
make_table(doc,
    ['Columna', 'Tipo', 'Descripción'],
    [
        ['id', 'uuid (PK)', 'Identificador único'],
        ['nombre', 'text', 'Nombre de quien pide oración'],
        ['apellidos', 'text', 'Apellidos (opcional)'],
        ['email', 'text', 'Correo electrónico (opcional)'],
        ['telefono', 'text', 'Teléfono (opcional)'],
        ['peticion', 'text', 'Texto completo de la petición de oración'],
        ['ubicacion', 'text', '¿De qué parte del mundo nos ves?'],
        ['referencia', 'text', '¿Cómo se enteró de la iglesia?'],
        ['decision_seguimiento', 'text', 'Respuesta al select: "si", "no" o "ya"'],
        ['created_at', 'timestamptz', 'Fecha y hora del envío'],
    ],
    col_widths=[4, 3, 9])

add_heading(doc, '7.9 Supabase Storage (Almacenamiento de imágenes)', 2)
add_body(doc,
    'Además de la base de datos, Supabase Storage almacena las imágenes del sitio. '
    'El bucket principal se llama "fotos" y es de acceso público (las URLs son '
    'accesibles directamente desde el navegador sin autenticación).')
add_body(doc,
    'Las URLs de Storage siguen el patrón:')
add_code(doc,
"""https://[project-id].supabase.co/storage/v1/object/public/fotos/[nombre-archivo.jpg]""")
add_note(doc,
    'No elimines imágenes directamente desde el panel de Supabase Storage si están referenciadas '
    'en la base de datos — usa siempre el panel admin del sitio, que elimina ambos registros '
    'en el mismo paso.', 'aviso')

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 9. GUÍA DE CAMBIOS COMUNES
# ═══════════════════════════════════════════════════════

add_heading(doc, '8. Guía de Cambios Comunes', 1)
add_body(doc,
    'Esta sección te guía paso a paso para realizar las modificaciones más frecuentes '
    'al sitio web, sin necesidad de conocimientos avanzados de programación.')

# 9.1
add_heading(doc, '8.1 Cambiar los colores del sitio', 2)
add_body(doc, 'El sitio usa 4 variables de color centrales. Cambiarlas es suficiente para '
    'transformar completamente la paleta del sitio.')
add_numbered(doc, 'Abre el archivo style.css.')
add_numbered(doc, 'Localiza el bloque :root al inicio del archivo.')
add_numbered(doc, 'Cambia los valores hexadecimales de las variables:')
add_code(doc,
""":root {
  --cream:  #F5EFE6;  /* Cambia este valor por el nuevo color crema    */
  --dark:   #1A1A1A;  /* Cambia este valor por el nuevo color oscuro    */
  --mid:    #888888;  /* Cambia este valor por el nuevo gris medio      */
  --border: rgba(26,26,26,0.15); /* Ajusta la opacidad si lo deseas    */
}""")
add_numbered(doc, 'Guarda el archivo y recarga el sitio en el navegador para ver los cambios.')
add_note(doc,
    'Puedes encontrar colores en hex en herramientas como coolors.co, colorhunt.co o '
    'simplemente buscando "color picker" en Google.', 'tip')

# 9.2
add_heading(doc, '8.2 Cambiar las fuentes del sitio', 2)
add_numbered(doc, 'Ve a fonts.google.com y elige la fuente que deseas.')
add_numbered(doc, 'Copia el enlace <link> que Google Fonts te proporciona.')
add_numbered(doc, 'En index.html, reemplaza el enlace de Google Fonts en el <head>:')
add_code(doc,
"""<!-- Reemplaza esta línea con el nuevo enlace de Google Fonts -->
<link href="https://fonts.googleapis.com/css2?family=NUEVA_FUENTE..." rel="stylesheet"/>""")
add_numbered(doc, 'En style.css, actualiza la propiedad font-family donde corresponda:')
add_code(doc,
"""/* Para cambiar la fuente del cuerpo (DM Sans) */
body {
  font-family: 'Nueva Fuente', sans-serif;
}

/* Para cambiar la fuente de títulos (Bebas Neue) */
.title, h1 {
  font-family: 'Otra Fuente', display;
}""")

# 9.3
add_heading(doc, '8.3 Agregar un nuevo horario', 2)
add_numbered(doc, 'Inicia sesión en el Panel Admin (admin.html).')
add_numbered(doc, 'Haz clic en la pestaña "Horarios" en el sidebar.')
add_numbered(doc, 'Llena el formulario con:')
add_bullet(doc, 'Título: nombre del servicio (ej: "Reunión de Damas")', 1)
add_bullet(doc, 'Día: el día de la semana (ej: "Viernes")', 1)
add_bullet(doc, 'Hora: la hora del servicio (ej: "6:00 PM")', 1)
add_numbered(doc, 'Haz clic en "Agregar Horario".')
add_numbered(doc, 'El horario aparecerá en el sitio público de inmediato.')
add_note(doc,
    'Si quieres controlar el orden en que aparecen los horarios, puedes editar '
    'directamente la columna "orden" en la tabla horarios desde el panel de Supabase.', 'nota')

# 9.4
add_heading(doc, '8.4 Publicar un anuncio', 2)
add_numbered(doc, 'Inicia sesión en el Panel Admin.')
add_numbered(doc, 'Ve a la pestaña "Anuncios".')
add_numbered(doc, 'Llena el formulario:')
add_bullet(doc, 'Título: título breve y claro del anuncio', 1)
add_bullet(doc, 'Fecha: la fecha relevante del anuncio', 1)
add_bullet(doc, 'Descripción: texto completo del anuncio', 1)
add_bullet(doc, 'Imagen: opcional — sube una foto si tienes', 1)
add_numbered(doc, 'Haz clic en "Publicar anuncio".')
add_numbered(doc, 'El anuncio aparecerá en el sitio (el sitio muestra los 3 más recientes activos).')
add_note(doc,
    'Si ya hay 3 anuncios activos publicados y publicas uno nuevo, el más antiguo '
    'dejará de verse en el sitio (aunque seguirá en la base de datos). '
    'Puedes desactivar los anuncios viejos manualmente con el botón de toggle.', 'aviso')

# 9.5
add_heading(doc, '8.5 Cambiar la foto de los pastores', 2)
add_numbered(doc, 'Inicia sesión en el Panel Admin.')
add_numbered(doc, 'Ve a la pestaña "Pastores".')
add_numbered(doc, 'Haz clic en "Subir foto" y selecciona la nueva imagen de los pastores.')
add_numbered(doc, 'Si la imagen es mayor de 2MB, el sistema la comprime automáticamente.')
add_numbered(doc, 'La nueva foto se activa de inmediato en el sitio público.')
add_note(doc,
    'Si subes varias fotos, el sistema mostrará la más reciente activa. '
    'Puedes desactivar las anteriores para mantener ordenado el archivo.', 'tip')

# 9.6
add_heading(doc, '8.6 Agregar una foto a la galería', 2)
add_numbered(doc, 'Inicia sesión en el Panel Admin.')
add_numbered(doc, 'Ve a la pestaña "Galería".')
add_numbered(doc, 'Haz clic en "Subir foto" y selecciona una o varias imágenes.')
add_numbered(doc, 'Las fotos se suben al Storage y quedan activas por defecto.')
add_numbered(doc, 'Aparecen en el sitio en el orden establecido.')
add_note(doc, 'Para cambiar el orden de las fotos, edita la columna "orden" directamente en Supabase.', 'nota')

# 9.7
add_heading(doc, '8.7 Cambiar el texto del hero', 2)
add_numbered(doc, 'Inicia sesión en el Panel Admin.')
add_numbered(doc, 'Ve a la pestaña "Hero".')
add_numbered(doc, 'Busca la sección de texto editable.')
add_numbered(doc, 'Modifica el subtítulo (lema de la iglesia) y/o la ubicación (ciudad).')
add_numbered(doc, 'Haz clic en "Guardar".')
add_numbered(doc, 'Los cambios se reflejan en el sitio al instante.')
add_note(doc,
    'El título grande "Linaje Escogido" está en el HTML directamente (elemento <h1>). '
    'Para cambiarlo, edita el archivo index.html.', 'nota')

# 9.8
add_heading(doc, '8.8 Agregar un nuevo campo al formulario Alpha', 2)
add_body(doc, 'Supongamos que queremos agregar un campo "¿Cuántos años tienes?". Sigue estos pasos:')
add_heading(doc, 'Paso 1 — Agregar la columna en Supabase', 3)
add_numbered(doc, 'Ve al panel de Supabase → Table Editor → registros_alpha.')
add_numbered(doc, 'Haz clic en "+ Add column".')
add_numbered(doc, 'Nómbrala "edad", tipo "text" (o "integer" si quieres solo números).')
add_numbered(doc, 'Guarda el cambio.')

add_heading(doc, 'Paso 2 — Agregar el campo en el HTML', 3)
add_body(doc, 'En index.html, dentro del formulario Alpha, agrega:')
add_code(doc,
"""<div class="field-wrap">
  <label class="field-label">¿Cuántos años tienes?</label>
  <input class="field-input" id="alphaEdad"
         type="number" placeholder="Tu edad" min="5" max="120"/>
</div>""")

add_heading(doc, 'Paso 3 — Capturar y guardar el dato en index.js', 3)
add_body(doc, 'En la función handleAlpha(), agrega la nueva variable y el campo al insert:')
add_code(doc,
"""async function handleAlpha(e) {
  e.preventDefault();
  const nombre    = document.getElementById('alphaNombre').value;
  const apellidos = document.getElementById('alphaApellidos').value;
  const email     = document.getElementById('alphaEmail').value;
  const telefono  = document.getElementById('alphaTelefono').value;
  const edad      = document.getElementById('alphaEdad').value; // ← NUEVO

  const { error } = await db
    .from('registros_alpha')
    .insert([{ nombre, apellidos, email, telefono, edad }]); // ← AGREGADO

  // ... resto de la función
}""")

# 9.9
add_heading(doc, '8.9 Agregar una nueva sección al sitio', 2)
add_body(doc, 'Ejemplo: agregar una sección "Ministerios" después de la sección de Pastores.')

add_heading(doc, 'Paso 1 — Agregar la sección en index.html', 3)
add_code(doc,
"""<section id="ministerios">
  <div class="label">Grupos de Servicio</div>
  <div class="title">Ministerios</div>
  <p>Descripción de los ministerios de la iglesia...</p>
  <!-- Aquí agrega el contenido que necesites -->
</section>""")

add_heading(doc, 'Paso 2 — Agregar el enlace en el nav', 3)
add_code(doc,
"""<ul class="nav-links">
  <!-- ... enlaces existentes ... -->
  <li><a href="#ministerios">Ministerios</a></li>
</ul>""")

add_heading(doc, 'Paso 3 — Estilizar en style.css (si es necesario)', 3)
add_code(doc,
"""#ministerios {
  padding: 80px 5%;
  background: var(--cream);
}
/* Agrega los estilos específicos que necesites */""")

add_note(doc,
    'Las clases .label y .title ya tienen estilos definidos globalmente en style.css, '
    'así que úsalas para mantener la consistencia visual de todas las secciones.', 'tip')

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 10. GLOSARIO
# ═══════════════════════════════════════════════════════

add_heading(doc, '9. Glosario de Términos Técnicos', 1)
add_body(doc,
    'Esta sección explica los términos técnicos usados en este manual, '
    'pensada para personas que están aprendiendo sobre desarrollo web.')

glossary = [
    ('API (Application Programming Interface)',
     'Una "interfaz" que permite que dos programas se comuniquen entre sí. '
     'En este proyecto, el JavaScript del sitio usa la API de Supabase para leer '
     'y escribir datos en la base de datos.'),
    ('async / await',
     'Palabras clave de JavaScript moderno para manejar operaciones que toman tiempo '
     '(como llamadas a internet). "async" declara que una función es asíncrona; '
     '"await" pausa la ejecución hasta que la operación termine. '
     'Ejemplo: await db.from("horarios").select() espera la respuesta de Supabase.'),
    ('Bucket',
     'En Supabase Storage, un "bucket" es como una carpeta raíz para organizar archivos. '
     'Este proyecto usa el bucket "fotos" para guardar todas las imágenes del sitio.'),
    ('CDN (Content Delivery Network)',
     'Red de servidores distribuidos globalmente que sirve archivos estáticos '
     '(como fuentes o librerías JavaScript) más rápido al usuario. '
     'Google Fonts y jsDelivr son CDNs usados en este proyecto.'),
    ('clamp()',
     'Función CSS que acepta tres valores: mínimo, preferido y máximo. '
     'Hace que un tamaño sea fluido entre esos límites según el ancho de pantalla. '
     'Ejemplo: font-size: clamp(32px, 5vw, 56px)'),
    ('CSS (Cascading Style Sheets)',
     'Lenguaje que controla la apariencia visual de una página web: colores, '
     'tipografía, espaciado, layout y animaciones. El archivo style.css contiene '
     'todos los estilos del proyecto.'),
    ('CSV (Comma-Separated Values)',
     'Formato de archivo de texto simple donde los datos se separan por comas. '
     'Se puede abrir con Excel o Google Sheets. El panel admin genera archivos CSV '
     'con los registros Alpha y peticiones de oración.'),
    ('DOM (Document Object Model)',
     'Representación en memoria del HTML de una página como un árbol de objetos. '
     'JavaScript puede manipular el DOM para cambiar el contenido visible '
     'sin recargar la página. Ejemplo: document.getElementById("horariosGrid").'),
    ('Flexbox',
     'Sistema de layout CSS que permite alinear y distribuir elementos '
     'en una fila o columna de forma flexible. Usado en la navegación, '
     'formularios y footer del sitio.'),
    ('Frontend',
     'La parte del sitio web que ve y con la que interactúa el usuario directamente '
     'en su navegador. En este proyecto: index.html, style.css e index.js.'),
    ('GitHub Pages',
     'Servicio gratuito de GitHub que sirve sitios web estáticos directamente '
     'desde un repositorio. Perfecto para este proyecto porque no requiere servidor.'),
    ('Grid (CSS Grid)',
     'Sistema de layout CSS bidimensional (filas y columnas). '
     'Se usa en las secciones de horarios, anuncios y galería para crear '
     'layouts de múltiples columnas que se adaptan al tamaño de pantalla.'),
    ('HTML (HyperText Markup Language)',
     'Lenguaje de marcado que define la estructura y contenido de una página web. '
     'Los archivos index.html y admin.html contienen la estructura del proyecto.'),
    ('JavaScript',
     'Lenguaje de programación que corre en el navegador y permite añadir '
     'interactividad y dinamismo a las páginas web. En este proyecto: index.js y admin.js.'),
    ('JSON (JavaScript Object Notation)',
     'Formato de intercambio de datos basado en texto, fácil de leer. '
     'Supabase devuelve los datos en formato JSON. '
     'Ejemplo: {"titulo": "Servicio Dominical", "hora": "10:00 AM"}'),
    ('Media Query',
     'Regla CSS que aplica estilos solo cuando se cumple una condición, '
     'como un tamaño de pantalla específico. '
     'Ejemplo: @media (max-width: 768px) { /* estilos para mobile */ }'),
    ('Opacity',
     'Propiedad CSS que controla la transparencia de un elemento. '
     'Va de 0 (invisible) a 1 (completamente visible). '
     'Se usa en el carrusel del hero para las transiciones entre imágenes.'),
    ('RLS (Row Level Security)',
     'Sistema de seguridad de Supabase/PostgreSQL que define reglas '
     'sobre qué usuarios pueden leer o modificar cada fila de cada tabla. '
     'Protege la base de datos incluso cuando la clave API es pública.'),
    ('Responsive Design',
     'Enfoque de diseño web donde el sitio se adapta automáticamente '
     'a diferentes tamaños de pantalla (móvil, tablet, escritorio) '
     'usando media queries, flexbox y CSS grid.'),
    ('Storage (Supabase)',
     'Servicio de almacenamiento de archivos de Supabase. '
     'Similar a Google Drive pero para aplicaciones. '
     'En este proyecto almacena todas las imágenes del sitio (fotos hero, galería, pastores, anuncios).'),
    ('Supabase',
     'Plataforma de backend como servicio (BaaS) que provee base de datos PostgreSQL, '
     'autenticación, storage y API REST automática. '
     'Es la alternativa open-source a Firebase de Google.'),
    ('Variable CSS (Custom Property)',
     'Valor reutilizable definido con --nombre: valor en CSS. '
     'Se usa con var(--nombre) en cualquier propiedad CSS. '
     'Cambiar la variable actualiza todos los elementos que la usan.'),
    ('Viewport',
     'El área visible de la página web en el navegador del usuario. '
     'En CSS, 100vh = altura completa del viewport, 100vw = ancho completo.'),
    ('z-index',
     'Propiedad CSS que controla el orden de apilamiento de elementos '
     'cuando se superponen. Mayor z-index = más al frente. '
     'Se usa para que el contenido del hero aparezca sobre las fotos de fondo.'),
]

for term, definition in glossary:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(2)
    run_term = para.add_run(term)
    run_term.font.bold = True
    run_term.font.name = 'Calibri'
    run_term.font.size = Pt(11)
    run_term.font.color.rgb = RGBColor(0x2A, 0x2A, 0x8A)

    para_def = doc.add_paragraph()
    para_def.paragraph_format.left_indent = Cm(0.7)
    para_def.paragraph_format.space_after = Pt(6)
    run_def = para_def.add_run(definition)
    run_def.font.name = 'Calibri'
    run_def.font.size = Pt(10.5)
    run_def.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# APÉNDICE: Referencia rápida
# ═══════════════════════════════════════════════════════

add_heading(doc, 'Apéndice — Referencia Rápida', 1)

add_heading(doc, 'A. Checklist para publicar el sitio', 2)
checks = [
    'Verificar que SUPABASE_URL y SUPABASE_KEY en index.js sean correctos.',
    'Confirmar que todas las tablas de Supabase tienen RLS configurado correctamente.',
    'Verificar que el bucket "fotos" en Supabase Storage es de acceso público.',
    'Revisar que el iframe de Google Maps tenga la URL correcta de la ubicación.',
    'Actualizar el año en el copyright del footer si es necesario.',
    'Reemplazar href="#" en el botón "Donar ahora" con la URL real de donaciones.',
    'Actualizar los enlaces de Facebook e Instagram con las URLs correctas.',
    'Verificar que los logos (logo.png, logo_inv.png, 1.png) estén en la raíz del proyecto.',
    'Probar el formulario Alpha y verificar que los registros lleguen a Supabase.',
    'Probar el formulario de Oración y verificar que las peticiones se guarden.',
    'Probar el panel admin y verificar que todas las pestañas funcionen.',
]
for check in checks:
    add_bullet(doc, check)

add_heading(doc, 'B. Resumen de IDs importantes del HTML', 2)
make_table(doc,
    ['ID del elemento', 'Qué es', 'Usado por'],
    [
        ['horariosGrid', 'Contenedor de las tarjetas de horario', 'loadHorarios() en index.js'],
        ['anunciosGrid', 'Contenedor de las tarjetas de anuncio', 'loadAnuncios() en index.js'],
        ['galeriaGrid', 'Contenedor de las fotos de galería', 'loadGaleria() en index.js'],
        ['alphaNombre', 'Input nombre del form Alpha', 'handleAlpha() en index.js'],
        ['alphaApellidos', 'Input apellidos del form Alpha', 'handleAlpha() en index.js'],
        ['alphaEmail', 'Input email del form Alpha', 'handleAlpha() en index.js'],
        ['alphaTelefono', 'Input teléfono del form Alpha', 'handleAlpha() en index.js'],
        ['horarios', 'Sección de horarios', 'Enlace del nav (#horarios)'],
        ['pastores', 'Sección de pastores', 'Enlace del nav (#pastores)'],
        ['anuncios', 'Sección de anuncios', 'Enlace del nav (#anuncios)'],
        ['ubicacion', 'Sección del mapa', 'Enlace del nav y botón hero'],
        ['alpha', 'Sección de formularios', 'Enlace del nav y botón hero'],
        ['oracion-form', 'Formulario de oración específicamente', 'Enlace del nav (#oracion)'],
    ],
    col_widths=[4, 5.5, 6.5])

add_heading(doc, 'C. Comandos útiles de Supabase (SQL)', 2)
add_body(doc, 'Si necesitas trabajar directamente en el SQL Editor de Supabase:')
add_code(doc,
"""-- Ver todos los horarios activos ordenados
SELECT * FROM horarios WHERE activo = true ORDER BY orden;

-- Ver los 3 anuncios más recientes activos
SELECT * FROM anuncios
WHERE activo = true
ORDER BY fecha DESC
LIMIT 3;

-- Ver todos los registros Alpha del mes actual
SELECT * FROM registros_alpha
WHERE fecha_registro >= date_trunc('month', now())
ORDER BY fecha_registro DESC;

-- Ver cuántas peticiones de oración hay en total
SELECT COUNT(*) FROM peticiones_oracion;

-- Desactivar todos los anuncios de golpe
UPDATE anuncios SET activo = false;

-- Cambiar el orden de un horario
UPDATE horarios SET orden = 1 WHERE titulo = 'Servicio Dominical';""")

add_heading(doc, 'D. Estructura de carpetas recomendada', 2)
add_code(doc,
"""Linaje_Escogido/
├── index.html          ← Página principal pública
├── style.css           ← Estilos globales
├── index.js            ← Lógica del sitio público
├── admin.html          ← Panel de administración
├── admin.js            ← Lógica del panel admin
├── logo.png            ← Logo oscuro (nav)
├── logo_inv.png        ← Logo claro (hero, footer)
├── 1.png               ← Favicon (ícono de pestaña)
├── pastores.jpg        ← Foto de respaldo de pastores
└── manual-codigo.docx  ← Este manual""")

# Página final
doc.add_page_break()
for _ in range(8):
    doc.add_paragraph()

final_para = doc.add_paragraph()
final_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
f_run = final_para.add_run('Linaje Escogido')
f_run.font.name = 'Calibri'
f_run.font.size = Pt(20)
f_run.font.bold = True
f_run.font.color.rgb = RGBColor(0x2A, 0x2A, 0x8A)

final_sub = doc.add_paragraph()
final_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
fs_run = final_sub.add_run('Una Iglesia Funcional · Apasionada · Sensible al Espíritu Santo')
fs_run.font.name = 'Calibri'
fs_run.font.size = Pt(12)
fs_run.font.italic = True
fs_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()
final_city = doc.add_paragraph()
final_city.alignment = WD_ALIGN_PARAGRAPH.CENTER
fc2_run = final_city.add_run('Reynosa, Tamaulipas · 2026')
fc2_run.font.name = 'Calibri'
fc2_run.font.size = Pt(11)
fc2_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

# ──────────────────────────────────────────────
# GUARDAR
# ──────────────────────────────────────────────
output_path = '/home/user/Linaje_Escogido/manual-codigo.docx'
doc.save(output_path)
print(f'Documento guardado exitosamente en: {output_path}')
